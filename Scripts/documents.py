import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from pdf2docx import Converter
import csv
import re

from _datetime import datetime


class DocumentFile():

    def __init__(self, path, convert=False):

        self.flags = []

        extension = 'docx'
        if '.' in path:
            extension = path.split('.')[1]
            path = path.split('.')[0]

        self.file_extension = extension
        self.path = path
        self.type = None
        self.expirations = {'product': None, 'sales_and_distribution': None, 'disposal_storage': None}
        if '-' in path:
            self.id = path.split('-')[1]
        else:
            self.id = path

        if convert:
            self.convert2docx()

        self.document_number = path.split('/')[-1][:8]

        self.crop_as_text = False

        self.load_exceptions()

    def __str__(self):
        return self.path

    def convert2docx(self):
        if self.file_extension == 'pdf':
            try:
                cv = Converter(f'{self.path}.{self.file_extension}')
                cv.convert(f'{self.path}.docx', start=0, end=None)
                cv.close()
            except:
                self.flags.append("Corrupted file. It was not possible to convert it to .docx")
            else:
                pass#os.remove(f'{self.path}.{self.file_extension}')
        else:
            result = os.system('unoconv -d document --format=docx "{path}.doc"'.format(path=self.path))
            if result == 0:
                os.remove("{path}.doc".format(path=self.path))
            else:
                self.flags.append("Corrupted file. It was not possible to convert it to .docx")

    def read_docx(self):
        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        doc = Document("{}.docx".format(self.path))
        if self.file_extension == 'pdf':
            styles = doc.styles
            styles.add_style('Normal + Left', WD_STYLE_TYPE.PARAGRAPH)
            for table in doc.tables:
                if not ('crops' in table.rows[0].cells[0].text.lower() or 'situation' in table.rows[0].cells[0].text.lower()):

                    texts = []
                    for i_row, row in enumerate(table.rows):
                        if '_' in row.cells[0].text and (i_row + 1) == len(table.rows) and\
                        '\n' not in row.cells[0].text:
                            texts.append(['ENDPAGE'])
                            continue

                        max_p = 0
                        for cell in row.cells:
                            max_p = max(max_p, len(cell.paragraphs))

                        for i in range(max_p):
                            text = []
                            for cell in row.cells:
                                try:
                                    if not '_' in cell.paragraphs[i].text:
                                        text.append(cell.paragraphs[i].text)
                                except:
                                    pass

                                if cell.tables:
                                    texts.append(["REMAINING TABLE", row.cells[0].tables])

                            if text:
                                texts.append(text)

                    for row in table.rows:
                        remove_row(table, row)

                    authorisation_ends = False
                    for i, text in enumerate(texts):
                        if text[0] == "REMAINING TABLE":
                            for table_ in text[1]:
                                table._element.addprevious(table_._tbl)
                        else:
                            new_text = text.copy()
                            deleted_text = ''
                            for j, t in enumerate(text):
                                try:
                                    if t.strip() == texts[i + 1][j].strip():
                                        new_text[j] = ''
                                        continue
                                    elif 'this authorisation ends' in t.lower():
                                        if authorisation_ends:
                                            new_text[j] = ''
                                            continue
                                        else:
                                            authorisation_ends = True
                                except:
                                    pass
                                try:
                                    if j > 0 and (t.strip() == new_text[j - 1].strip() or t.strip() == deleted_text.strip()):
                                        deleted_text = new_text[j]
                                        new_text[j] = ''
                                        continue
                                except:
                                    pass
                                deleted_text = ''
                            if '_' in text[0]:
                                continue
                            paragraph = doc.add_paragraph('\t'.join(new_text), style='Normal + Left')
                            table._element.addprevious(paragraph._p)

            doc.save("{}.docx".format(self.path))

        return doc

    def load_exceptions(self):

        # Pesticide exceptions
        try:
            reader = csv.reader(open('../../../../exception_pesticides.csv', 'r'))
        except:
            try:
                reader = csv.reader(open('../../../exception_pesticides.csv', 'r'))
            except:
                reader = csv.reader(open('exception_pesticides.csv', 'r'))

        self.pesticides_exceptions = []
        for row in reader:
            self.pesticides_exceptions.append(row[0].lower())

        # Extensions exceptions
        try:
            reader = csv.reader(open('../../../../exception_extensions.csv', 'r'))
        except:
            try:
                reader = csv.reader(open('../../../exception_extensions.csv', 'r'))
            except:
                reader = csv.reader(open('exception_extensions.csv', 'r'))
        self.extensions_exceptions = {}
        for row in reader:
            self.extensions_exceptions[row[0].lower()] = row[1:]

    def check_date(self, date_str):

        if not date_str:
            self.flags.append("A value for date was not found")

            return None

        new_date_str = date_str.replace(' ', '').replace('for', '').replace('except', '').strip()
        possible_formats = ["%d%B%Y", "%dst%B%Y", "%dth%B%Y", "%d/%m/%Y"]

        for f in possible_formats:
            try:
                date_obj = datetime.strptime(new_date_str, f)

                return date_obj.strftime("%d/%m/%Y")
            except ValueError:
                continue

        if 'granted a period' in date_str:
            self.flags.append("A value for date was not found")
            return None
        self.flags.append("The following date is not in the right format: {}".format(date_str))

        return 'Invalid Date'

    def get_main_info(self, doc):
        '''
        Document ID
        Date of Issue
        Document Type
        Approval Number
        Product Expiry
        Expiry (sale and distribution of existing stocks)
        Expiry (disposal storage and use of existing stocks)

        :param doc:
        :return:
        '''

        # Document Type
        try:
            if 'Authorisation Number:' in doc.paragraphs[0].text or 'Authorisation Number:' in doc.paragraphs[1].text or 'Approval Number' in doc.paragraphs[0].text or 'Approval Number' in doc.paragraphs[1].text:
                self.type = 'Authorization'
        except:
            pass

        # Exceptions
        expirations = []

        # Flag exceptions
        key_phrases = ['use(s) to be withdrawn', 'use to be withdrawn', 'revocation of use on certain crops', 'revocation of use on']
        found_key_phrases = []

        # Multiple components formulation
        extra_formulation = False

        for i, p in enumerate(doc.paragraphs):
            #print(p.text)
            # Document Type
            if 'Amendment Notice' in p.text or 'Label amendments:' in p.text or p.text == 'Amendment Table' or p.text == 'THE AMENDMENT':
                self.type = 'Amendment'
            elif 'Withdrawal Notice' in p.text:
                self.type = 'Withdrawal'
            elif 'parallel trade' in p.text.lower():
                self.type = 'Parallel'
            elif p.text == 'NOTICE OF APPROVAL':
                self.type = 'Authorization'

            # Expiry
            if (p.text == 'This authorisation ends:' or p.text == 'This approval ends:') and not expirations:
                expirations = [i + 1, i + 2, i + 3]
            elif ('This authorisation ends:' in p.text and 'regulatory' not in p.text) and not expirations:
                expirations = [i, i + 1, i + 2]
                break
            elif ('This approval ends' in p.text) and not expirations:
                expirations = [i + 1, i + 2, i + 3]


            # Get Date of Issue
            if 'date of issue' in p.text.lower():
                p.text = p.text.strip().lower()

                '''
                if '\t' in p.text:
                    pos = p.text.find('\t') # Change from find to rfind for 12726
                else:
                    pos = p.text.find(' ', 10)
                '''

                p.text = p.text.replace('\t', '').replace('date of issue', '').replace(':', '').strip()
                self.date_of_issue = p.text

                if ')' in self.date_of_issue:
                    pos = self.date_of_issue.find(')')
                    self.date_of_issue = self.date_of_issue[pos + 1:]

                if '.' in self.date_of_issue:
                    pos = self.date_of_issue.rfind('.')
                    self.date_of_issue = self.date_of_issue[:pos]

                if '.' in self.date_of_issue:
                    pos = self.date_of_issue.find('.')
                    self.date_of_issue = self.date_of_issue[pos:]

                self.date_of_issue = self.date_of_issue.strip()

            # Formulation
            if extra_formulation and '\t' in p.text:
                extra_formulation = False
            elif 'Formulation' in p.text or extra_formulation and p.text:
                if not hasattr(self, 'formulation'):
                    self.formulation = p.text.replace('Formulation:', '').strip()
                else:
                    self.formulation += '\n' + p.text.strip()
                extra_formulation = True


            # Extent of authorization
            if 'Extent of Authorisation:' in p.text:
                self.extent_of_authorization = p.text.replace('Extent of Authorisation:', '').strip()

            # Flag exceptions
            for key_phrase in key_phrases:
                if key_phrase in p.text.lower():
                    found_key_phrases.append(key_phrase)
                    break

            if 'Maximum individual dose:' in p.text or 'situations:' in p.text.lower():
                self.crop_as_text = True

        for table in doc.tables:
            if table.rows:
                headers = ' '.join([cell.text.lower() for cell in table.rows[0].cells])
                for key_phrase in key_phrases:
                    if key_phrase in headers:
                        found_key_phrases.append(key_phrase)
                        break

        if found_key_phrases:
            for phrase in found_key_phrases:
                self.flags.append('Key phrase "{}" found on {}'.format(phrase, self.path.split('/')[-1]))

        if not self.type:
            if '/' in self.path:
                type = self.path.split('/')[-1]
            else:
                type = self.path
            type = type.split('-')
            if len(type) >= 3:
                type = type[2]
            else:
                type = 'Approval'

            if type == 'Approval':
                self.type = 'Authorization'

        if hasattr(self, 'date_of_issue'):
            self.date_of_issue = self.check_date(self.date_of_issue)
        # Get expiration dates
        if expirations:
            if self.file_extension == 'pdf':
                expirations.append(expirations[-1] + 1)
                expirations.append(expirations[-1] + 2)
                expirations.append(expirations[-1] + 3)
            for e in expirations:
                date = doc.paragraphs[e].text.strip()
                if 'This authorisation remains' in date:
                    continue

                if 'except' in date:
                    key = 'product'
                elif 'distribution' in date:
                    key = 'sales_and_distribution'
                elif 'storage' in date:
                    key = 'disposal_storage'
                else:
                    continue

                if "\t" in date or '   ' in date or self.file_extension == 'pdf':
                    if ')' in date:
                        date = date[date.find(')') + 2:].strip()
                    else:
                        date = date[date.find('\t') + 1:].strip()

                if date.find('/') <  date.find(' ') and date.find('/') > 0:
                    date = date.split(' ')[0]
                else:
                    date = date.split(' ')[:3]
                self.expirations[key] = self.check_date(' '.join(date))

    def get_protections(self, doc):

        protections = {'Environmental protection:':{'headers':[]},
                       'Operator protection:':{'headers':[]},
                      'Other specific restrictions:':{'headers':[]}}

        key = None
        header = None
        header_intent = None
        ps = []
        if self.file_extension == 'pdf':
            for i, p in enumerate(doc.paragraphs):
                if not p.text:
                    continue
                elif p.text == 'ENDPAGE':
                    last_line = p.text
                    continue

                found = False
                for k in protections.keys():
                    if k in p.text.capitalize():
                        found = True
                        key = k
                        break
                if found:
                    # key = p.text.capitalize()
                    first = True
                    if header:
                        header_key = header['key']
                        del header['key']
                        protections[header_key]['headers'].append(header)
                        header = None

                    if not p.text.capitalize().replace(key, '').strip():
                        continue
                    else:
                        p.text = p.text.lower().replace(key.lower(), '')
                elif re.findall(r'(^|\s)\(*(\d|\w)\)', p.text.lower()) and header:
                    key = header['key']
                elif header and last_line == 'ENDPAGE':
                    pass
                elif p.text == '' or p.text == '\n' or p.text == '\t' or (
                        p.style.name != 'Restriction Number' and p.style.name != 'Restriction Letter' and 'Normal + Left' not in p.style.name):
                    key = None

                if key:

                    ps.append(p)
                    p.text = p.text.replace('\n', '').replace('\t', ' ').strip()

                    if p.paragraph_format.first_line_indent == None and p.style.name == 'Restriction Number' or first:
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                        header_intent = p.paragraph_format.first_line_indent
                        first = False
                    elif p.paragraph_format.first_line_indent == 0 and p.style.name == 'Restriction Number':
                        if not header:
                            header = {'value': p.text, 'subheaders': [], 'key': key}
                        else:
                            header['value'] += '\n' + p.text
                    elif re.findall(r'(^|\s)\(*([a-z])\)', p.text.lower()) and header:
                        header['subheaders'].append(p.text)
                    elif re.findall(r'(^|\s)\(*([0-9])\)', p.text.lower()):
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                    elif last_line == 'ENDPAGE' and header:
                        header['value'] += ' ' + p.text
                    elif p.paragraph_format.first_line_indent == None:
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                        header_intent = p.paragraph_format.first_line_indent
                    elif (p.style.name == 'Restriction Letter' or 'Normal + Left' in p.style.name or (
                            p.paragraph_format.first_line_indent < 0 and p.paragraph_format.first_line_indent != header_intent)) and header:
                        header['subheaders'].append(p.text)
                    else:
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                        header_intent = p.paragraph_format.first_line_indent

                last_line = p.text
        else:
            for i, p in enumerate(doc.paragraphs):
                if not p.text:
                    continue
                elif p.text == 'ENDPAGE':
                    last_line = p.text
                    continue

                found = False
                for k in protections.keys():
                    if k in p.text.capitalize():
                        found = True
                        key = k
                        break
                if found:
                    # key = p.text.capitalize()
                    first = True
                    if header:
                        header_key = header['key']
                        del header['key']
                        protections[header_key]['headers'].append(header)
                        header = None

                    if not p.text.capitalize().replace(key, '').strip():
                        continue
                    else:
                        p.text = p.text.lower().replace(key.lower(), '')
                elif re.findall(r'(^|\s)\(*(\d|\w)\)', p.text.lower()) and header:
                    key = header['key']
                elif header and last_line == 'ENDPAGE':
                    pass
                elif p.text == '' or p.text == '\n' or p.text == '\t' or (
                        p.style.name != 'Restriction Number' and p.style.name != 'Restriction Letter' and 'Normal + Left' not in p.style.name):
                    key = None

                if key:

                    ps.append(p)
                    p.text = p.text.replace('\n', '').replace('\t', ' ').strip()

                    if p.paragraph_format.first_line_indent == None and p.style.name == 'Restriction Number' or first:
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                        header_intent = p.paragraph_format.first_line_indent
                        first = False

                    elif (re.findall(r'(^|\s)\(*([a-z])\)', p.text.lower()) or p.style.name == 'Restriction Letter') and header:
                        header['subheaders'].append(p.text)

                    elif (p.paragraph_format.first_line_indent == 0 or p._p.pPr.ind) and p.style.name == 'Restriction Number':
                        if not header:
                            header = {'value': p.text, 'subheaders': [], 'key': key}
                        else:
                            header['value'] += '\n' + p.text

                    elif re.findall(r'(^|\s)\(*([0-9])\)', p.text.lower()):
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                    elif last_line == 'ENDPAGE' and header:
                        header['value'] += ' ' + p.text
                    elif p.paragraph_format.first_line_indent == None:
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                        header_intent = p.paragraph_format.first_line_indent
                    elif (p.style.name == 'Restriction Letter' or 'Normal + Left' in p.style.name or (
                            p.paragraph_format.first_line_indent < 0 and p.paragraph_format.first_line_indent != header_intent)) and header:
                        header['subheaders'].append(p.text)
                    else:
                        if header:
                            header_key = header['key']
                            del header['key']
                            protections[header_key]['headers'].append(header)
                        header = {'value': p.text, 'subheaders': [], 'key': key}
                        header_intent = p.paragraph_format.first_line_indent

                last_line = p.text

        if header:
            header_key = header['key']
            del header['key']
            protections[header_key]['headers'].append(header)

        for k, v in protections.items():
            if not protections[k]['headers']:
                protections[k]['headers'].append({'value':'', 'subheaders':[]})

        return protections

    def get_crops(self, doc):

        def fix_parenthesis(text):
            count_open_par = text.count("(")
            count_close_par = text.count(")")

            if count_close_par or count_close_par:
                opens = []
                closes = []
                last_open = 0
                for i, c in enumerate(text):
                    if c == '(':
                        opens.append(i)
                        last_open = i
                    elif c == ')':
                        if opens:
                            del opens[-1]
                        else:
                            closes.append(i)

                if closes:
                    closes.reverse()
                    for ind in closes:
                        text = text[:ind] + text[ind + 1:]

                if opens:
                    opens.reverse()
                    for ind in opens:
                        if ind == last_open:
                            text += ')'
                        else:
                            text = text[:ind] + text[ind + 1:]

                return text.strip()

            if count_open_par > count_close_par:
                pos = text.find("(")
                text = text[:pos] + text[pos + 1:]
            elif count_open_par < count_close_par:
                pos = text.rfind(")")
                text = text[:pos] + text[pos + 1:]

            return text.strip()

        def get_metric(text):
            if '(' in text:
                #text = text[text.find('(')+1:]
                text = fix_parenthesis(text)
                new_text = re.findall(r'\((.+)\)', text)

                if new_text:
                    return new_text[0]
                else:
                    text = text[text.find('(') + 1:]
                    text = fix_parenthesis(text)
                    return text
            return None

        def get_crops(text):

            if self.file_extension == 'pdf':
                text = text.replace('\n', ' ')

            crops = []

            for excep, new_value in self.extensions_exceptions.items():
                if excep in text.lower():
                    #print("Exception extension found for documents.py:", text)
                    text = text.lower().replace(excep, '')
                    crops += new_value

            for excep in self.pesticides_exceptions:
                if excep in text.lower():
                    #print("Exception pesticide found for documents.py:", text)
                    text = text.lower().replace(excep, '')
                    crops.append(excep.capitalize())

            text = fix_parenthesis(text)

            if not text:
                return crops, []

            additional_names = []

            while '\n\n\n' in text:
                text = text.replace('\n\n\n', '\n\n')

            text = text.replace('\nor', ',').replace('\nOR', ',').replace('or\n', ',').replace('OR\n', ',')

            if self.file_extension == 'pdf':
                text = text.replace('\n\n', ' ')
                text = text.replace('\n', ' ')
            else:
                text = text.replace('\n\n', ',')

            ini = 0
            pause = False
            for i, c in enumerate(text):
                if c == ',' and not pause:
                    crop_name = text[ini:i].strip().capitalize()
                    if crop_name:
                        if crop_name.lower() not in ['or', '.', '']:
                            crops.append(crop_name)
                    ini = i + 1
                elif text[i:i+5] == ' and ' and not pause:
                    crop_name = text[ini:i].strip().capitalize()
                    if crop_name:
                        if crop_name.lower() not in ['or', '.', '']:
                            crops.append(crop_name)
                    ini = i + 5
                elif c == '(':
                    pause = True
                elif c == ')':
                    pause = False

            if text[ini:].strip():
                if 'a)' in text[ini:2] or ' a)' in text[ini:] or \
                        'b)' in text[ini:2] or ' b)' in text[ini:] or \
                        'c)' in text[ini:2] or ' c)' in text[ini:]:
                    additional_names.append(text[ini:].strip().capitalize())
                elif text[ini:].strip().lower() not in ['or', '.']:
                    crops.append(text[ini:].strip().capitalize())

            return crops, additional_names

        def find_crop_tables(doc):

            for table in doc.tables:
                if table.rows:
                    if 'crops' in table.rows[0].cells[0].text.lower() or 'situation' in table.rows[0].cells[0].text.lower():
                        if 'aquatic buffer' in table.rows[0].cells[1].text.lower():
                            aquatic_tables.append(table)
                        else:
                            crops_tables.append(table)
                    elif table.rows[0].cells[0].tables:
                        new_doc = table.rows[0].cells[0]
                        for row in table.rows[1:]:
                            new_row = new_doc.tables[0].add_row()
                            for i, cell in enumerate(row.cells):
                                new_row.cells[i].text = cell.text
                        find_crop_tables(new_doc)

        crops_tables = []
        aquatic_tables = []
        if doc.tables:

            find_crop_tables(doc)

            #if not crops_tables:
            #    return None, None

        """
        # This section was made to obtain all possible crops value
        if crops_tables:
            all_crops = []
            for crops_table in crops_tables:

                mid_ind = None
                header = ''
                headers = {}
                for i, cell in enumerate(crops_table.rows[0].cells):
                    if 'maximum individual dose' in cell.text.lower().replace('\n', ''):
                        header = cell.text
                        mid_ind = i
                        headers['mid'] = {}
                        headers['mid']['header'] = cell.text
                        headers['mid']['index'] = i
                    elif 'maximum total dose' in cell.text.lower().replace('\n', ''):
                        headers['mtd'] = {}
                        headers['mtd']['header'] = cell.text
                        headers['mtd']['index'] = i
                    elif 'maximum number of treatments' in cell.text.lower().replace('\n', ''):
                        headers['mnt'] = {}
                        headers['mnt']['header'] = cell.text
                        headers['mnt']['index'] = i
                    elif 'latest time of application' in cell.text.lower().replace('\n', ''):
                        headers['lta'] = {}
                        headers['lta']['header'] = cell.text
                        headers['lta']['index'] = i

                for row in crops_table.rows[1:]:
                    keys = ['mid', 'mtd', 'mnt', 'lta']
                    crop = [row.cells[0].text]
                    for key in keys:
                        if headers.get(key):
                            crop += [headers[key]['header'], row.cells[headers[key]['index']].text]
                    crop.append('/'.join(self.path.split('/')[-2:]))
                    all_crops.append(crop)

            return all_crops, None
        """

        if crops_tables:
            # CROPS
            crops_result = {}

            for crops_table in crops_tables:
                row_crops = []
                keys = {'crop': 0}
                metrics = {}
                for i, cell in enumerate(crops_table.rows[0].cells):
                    cell.text = cell.text.replace('\n', '')
                    if 'maximum individual dose' in cell.text.lower():
                        key = 'mid'
                    elif 'maximum total dose' in cell.text.lower():
                        key = 'mtd'
                    elif 'maximum number of treatments' in cell.text.lower():
                        key = 'mnt'
                    elif 'latest time of application' in cell.text.replace('\n', ' ').lower():
                        key = 'lta'
                    elif 'maximum' in cell.text.lower() and 'total' in cell.text.lower() and 'dose' in cell.text.lower():
                        key = 'mtd'
                        cell.text = cell.text.replace('maximum', '').replace('total', '').replace('dose',
                                                                                                       '').strip()
                    elif 'maximum' in cell.text.lower() and 'individual' in cell.text.lower() and \
                            'dose' in cell.text.lower():
                        key = 'mid'
                        cell.text = cell.text.replace('maximum', '').replace('individual', '').replace('dose', '').strip()
                    elif 'maximum' in cell.text.lower() and 'number' in cell.text.lower() and 'treatments' in cell.text.lower():
                        key = 'mnt'
                        cell.text = cell.text.replace('maximum', '').replace('number', '').replace('treatments',
                                                                                                       '').strip()
                    elif 'latest' in cell.text.lower() and 'time' in cell.text.lower() and 'application' in cell.text.lower():
                        key = 'lta'
                        cell.text = cell.text.replace('latest', '').replace('time', '').replace('application',
                                                                                                       '').strip()
                    else:
                        continue

                    keys[key] = i
                    metric = get_metric(cell.text)
                    if metric:
                        metrics[key] = metric

                if len(keys) == 1:
                    continue

                available_cols = ['mid', 'mtd', 'mnt', 'lta']
                stage = 'primary'
                for row in crops_table.rows[1:]:
                    row.cells[1].text = row.cells[1].text.replace('\n', '')
                    if 'and/or' in row.cells[1].text.strip().lower():
                        stage = 'and_or'
                        continue
                    elif 'AND' in row.cells[1].text.strip() or 'and:' == row.cells[1].text.strip().lower():
                        stage = 'and'
                        continue
                    elif 'either:' in row.cells[1].text.strip().lower() or 'or' == row.cells[1].text.strip().lower() or 'or:' == row.cells[1].text.strip().lower():
                        stage = 'or'
                        continue
                    elif len(row.cells) == 2:
                        if (not row.cells[0].text.strip() or row.cells[0].text.strip() == '-') and (not row.cells[1].text.strip() or row.cells[1].text.strip() == '-'):
                            continue
                    elif len(row.cells) > 2:
                        if (not row.cells[0].text.strip() or row.cells[0].text.strip() == '-') and (not row.cells[1].text.strip() or row.cells[1].text.strip() == '-') and (not row.cells[2].text.strip() or row.cells[2].text.strip() == '-'):
                            continue

                    text_cell = row.cells[keys['crop']].text

                    values = [{"crop": ""}]
                    additional_names = []
                    for key in available_cols:
                        if keys.get(key):
                            value = row.cells[keys[key]].text.strip()
                            if self.file_extension == 'pdf':
                                value = value.replace('\n', '')

                            while '  ' in value:
                                value = value.replace('  ', ' ')

                            if value[:2] == 'i)' or value[:3] == 'ii)' or value[:4] == 'iii)':
                                value = '(' + value
                            value = fix_parenthesis(value)
                            if '(i)' in value or '(ii)' in value or '(iii)' in value or 'and/or' in value:
                                if 'and/or' in value:
                                    for sep in ['(iii)', '(ii)', '(i)']:
                                        value = value.replace(sep, '').strip()
                                    multiple_values = value.split('and/or')
                                    multiple_values.reverse()
                                else:
                                    multiple_values = []
                                    for sep in ['(iii)', '(ii)', '(i)']:
                                        if sep in value:
                                            value_ = value[value.find(sep) + len(sep):].strip()
                                            if ':' in value:
                                                additional_names.append(value_[:value_.find(':')].strip())
                                                value_ = value_[value_.find(':') + 1:].strip()
                                            elif '(weed control)' in value:
                                                additional_names.append('Weed control')
                                                value_ = value_.replace('(weed control)', '')
                                            elif 'Weed control' in value:
                                                additional_names.append('Weed control')
                                                value_ = value_.replace('Weed control', '')
                                            elif 'Chemical thinning (by injection)' in value:
                                                additional_names.append('Chemical thinning (by injection)')
                                                value_ = value_.replace('Chemical thinning (by injection)', '')
                                            elif 'Chemical thinning' in value:
                                                additional_names.append('Chemical thinning')
                                                value_ = value_.replace('Chemical thinning', '')
                                            elif 'Stump application' in value:
                                                additional_names.append('Stump application')
                                                value_ = value_.replace('Stump application', '')

                                            multiple_values.append(value_)
                                            value = value[:value.find(sep)]
                            elif '\n\n' in value:
                                while '\n\n\n' in value:
                                    value = value.replace('\n\n\n', '\n\n')

                                multiple_values = value.split('\n\n')
                            else:
                                if ':' in value and key != 'mnt':
                                    if value.find('(') == -1 or (value.find('(') > value.find(':')):
                                        if 'Maximum total dose' not in value[:value.find(':')].strip():
                                            if value[:value.find(':')].strip().lower() not in ['or', 'and']:
                                                additional_names.append(value[:value.find(':')].strip())
                                                value = value[value.find(':') + 1:].strip()
                                elif '(weed control)' in value:
                                    additional_names.append('Weed control')
                                    value = value.replace('(weed control)', '')
                                elif 'Weed control' in value:
                                    additional_names.append('Weed control')
                                    value = value.replace('Weed control', '')
                                elif 'Chemical thining (by injection)' in value:
                                    additional_names.append('Chemical thining (by injection)')
                                    value = value.replace('Chemical thining (by injection)', '')
                                elif 'Chemical thinning' in value:
                                    additional_names.append('Chemical thinning')
                                    value = value.replace('Chemical thinning', '')

                                multiple_values = [value]

                            for vi, value in enumerate(multiple_values):

                                if len(values)-1 < vi:
                                    values.append(values[vi-1].copy())

                                value = value.replace('\xa0', ' ').replace('pre-crop emergence', '').\
                                    replace('1l ', '1 l ').replace('2l ', '2 l ').replace('max. ', '')\
                                    .replace('0.2 0.2 litres', '0.2 litres')\
                                    .replace('0.2 * 0.2 litres', '0.2 litres').strip()

                                values[vi][key] = value
                                if metrics.get(key):
                                    values[vi][key + '_metric'] = metrics[key]

                                '''
                                if value in ['- See ‘other specific restrictions’.', '(see ‘Other Specific Restrictions’)', '(see ‘Other specific restrictions’)', 'See Other Specific Restrictions', 'See Other specific restriction (1)', 'See Other Specific Restriction 2 (applied as a spray/drench)', '100 added to 1 litre in a hand-held sprayer']:
                                    value = '-'
                                if metrics.get(key) and ('/' not in value and 'per' not in value and 'see ' not in value.lower()):
                                    values[vi][key + '_metric'] = metrics[key]
                                    values[vi][key] = value
                                elif (' ' in value or '/' in value or 'per' in value) and key != 'lta':
                                    if 'see ' in value.lower() and '(see ' not in value.lower() and value.lower().find('see ') < 13:
                                        values[vi][key] = value
                                    elif ' ' in value and ('see ' in value.lower() or 'ocr' in value.lower() or 'specific restriction' in value.lower()):
                                        value = value.split(' ')
                                        values[vi][key] = value[0]
                                        if metrics.get(key):
                                            values[vi][key + '_metric'] = metrics[key]
                                    elif ' ' in value:
                                        value = value.split(' ')
                                        values[vi][key + '_metric'] = ' '.join(value[1:])
                                        values[vi][key] = value[0]
                                    elif '/' in value:
                                        ini_pos = value.find('/')
                                        ipos = 0
                                        for ipos in range(ini_pos, 0, -1):
                                            if value[ipos] in ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0']:
                                                break

                                        values[vi][key + '_metric'] = value[ipos + 1:]
                                        values[vi][key] = value[:ipos + 1]
                                    elif 'per' in value:
                                        ini_pos = value.find('per')
                                        ipos = 0
                                        for ipos in range(ini_pos, 0, -1):
                                            if value[ipos] in ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0']:
                                                break

                                        values[vi][key + '_metric'] = value[ipos + 1:]
                                        values[vi][key] = value[:ipos + 1]
                                elif value:
                                    values[vi][key] = value
                                    if key != 'lta' and value != '-':
                                        self.flags.append('no metric found for {} (in crops) with value: {}'.format(key, value))
                                '''

                    new_crops, additional_names2 = get_crops(text_cell)
                    additional_names += additional_names2
                    if new_crops:
                        row_crops = new_crops

                    hard_replaces = {'(a)':'', '(b)':'', 'Admixture: ':'', '\t':'', '<i>':'', '</I>':''}
                    for c in row_crops:
                        for vi, values_ in enumerate(values):
                            val = values_.copy()
                            val['crop'] = c
                            if len(additional_names) > vi:
                                val['crop'] += ' - ' + additional_names[vi]

                            for k_, v_ in hard_replaces.items():
                                val['crop'] = val['crop'].replace(k_, v_).strip()

                            if val['crop'] and (val.get('mid') or val.get('mnt') or val.get('lta')):
                                if val['crop'] in ['', 'or', '.']:
                                    continue
                                if not val['crop'] in crops_result:
                                    stage = 'primary'
                                    crops_result[val['crop']] = {}
                                elif stage == 'primary':
                                    stage = 'OR'

                                if stage not in crops_result[val['crop']]:
                                    crops_result[val['crop']][stage] = []

                                crops_result[val['crop']][stage].append(val)

            del row_crops

            # AQUATIC BUFFER
            row_crops = []
            keys = {'crop': 0}
            metrics = {}
            aquatic_result = []
            for aquatic_table in aquatic_tables:
                for i, cell in enumerate(aquatic_table.rows[0].cells):
                    if 'distance' in cell.text.lower():
                        key = 'dis'
                    else:
                        continue

                    keys[key] = i
                    metric = get_metric(cell.text)
                    if metric:
                        metrics[key] = metric

                available_cols = ['dis']
                for row in aquatic_table.rows[1:]:
                    if 'OR' in row.cells[1].text.strip():
                        continue

                    text_cell = row.cells[keys['crop']].text

                    values = {"crop": ""}
                    for key in available_cols:
                        if keys.get(key):
                            value = fix_parenthesis(row.cells[keys[key]].text.strip())

                            if metrics.get(key):
                                values[key + '_metric'] = metrics[key]
                                values[key] = value
                            elif ' ' in value and key != 'lta':
                                value = value.split(' ')
                                values[key + '_metric'] = ' '.join(value[1:])
                                values[key] = value[0]
                            else:
                                values[key] = value
                                if key != 'lta' and value != '-':
                                    self.flags.append('no metric found for {} (in aquatic) with value: {}'.format(key, value))

                    new_crops, _ = get_crops(text_cell)
                    if new_crops:
                        row_crops = new_crops

                    for c in row_crops:
                        val = values.copy()
                        val['crop'] = c
                        aquatic_result.append(val)

            if not aquatic_result:
                aquatic_result = None

            return crops_result, aquatic_result
        elif self.crop_as_text:
            crops = []
            crop_names = []
            new_crop = {}
            fields = {'crop':['crops:', 'situations:'], 'mid':['Maximum individual dose:'], 'mtd':['Maximum total dose:'], 'lta':['Latest time of application:']}
            for i, p in enumerate(doc.paragraphs):
                if not p.text:
                    continue
                for name, field_vals in fields.items():
                    found = False
                    for field in field_vals:
                        if field in p.text.lower() and name == 'crop':
                            new_crop = {}
                            crop_names, _ = get_crops(p.text[p.text.find(':') + 1 :].strip())
                            new_crop['crop'] = crop_names[0]
                            found = True
                            break
                        elif field in p.text and new_crop:
                            if name in new_crop:
                                for crop_name in crop_names:
                                    new_crop_copy = new_crop.copy()
                                    new_crop_copy['name'] = crop_name
                                    crops.append(new_crop_copy)
                                new_crop = {}
                                crop_names = []
                            value = p.text.replace(field, '').strip()
                            if ' ' in value and name != 'lta':
                                if value == 'See label':
                                    new_crop[name] = value
                                else:
                                    value = value.split(' ')
                                    new_crop[name + '_metric'] = ' '.join(value[1:])
                                    new_crop[name] = value[0]
                            else:
                                new_crop[name] = value

                            found = True
                            break
                    if found:
                        break

            if crop_names:
                for crop_name in crop_names:
                    new_crop_copy = new_crop.copy()
                    new_crop_copy['crop'] = crop_name
                    crops.append(new_crop_copy)

            return crops, None
        else:
            return [], None

    def get_parallel_parent(self, doc):
        for i, p in enumerate(doc.paragraphs):
            if 'UK reference product:' in p.text:
                parent = p.text[p.text.find('UK reference product:') + 21 :]
                parent_name = parent[: parent.find('(')].strip()
                parent_mapp = parent[parent.find('MAPP No ') + 8:-1]

                return {'name': parent_name, 'mapp': parent_mapp}
            elif 'UK reference product' in p.text:
                parent = p.text[p.text.find('UK reference product') + 21:]
                parent_name = parent[: parent.find('(')].strip().replace("‘", "").replace("’", "")
                parent_mapp = parent[parent.find('MAPP No ') + 8 : parent.find(")")]

                return {'name':parent_name.strip(), 'mapp':parent_mapp}
        return {'name': 'Not found', 'mapp': 'Not found'}

    def get_active_substances(self, doc):
        substances = []

        for p in doc.paragraphs:
            if 'Amendment Notice' in p.text:
                if ' – ' in p.text:
                    name = p.text.split(' – ')[-1]
                elif ' - ' in p.text:
                    name = p.text.split(' - ')[-1]
                else:
                    continue

                if ' – ' in name:
                    name = name.split(' – ')[-1]
                elif ' - ' in name:
                    name = name.split(' - ')[-1]
                elif ':' in name:
                    name = name.split(':')[-1]

                if 'extension to expiry date' in name.lower() or 'amendment of operator protection'  in name.lower():
                    continue

                names = [name]
                separators = [' and ', ' or ', ', ', '.']
                for s in separators:
                    new_names = []
                    for n in names:
                        if n:
                            new_names += n.split(s)
                    names = new_names

                for name in names:
                    if name.strip().capitalize():
                        substance = {}
                        substance['name'] = name.strip().capitalize()
                        substance['expirations'] = self.expirations
                        substance['documents'] = [{'date_of_issue':self.date_of_issue if hasattr(self, 'date_of_issue') else None,
                                                   'type': self.type,
                                                   'document_id': self.id}]
                        substances.append(substance)

                break

        return substances

    def is_empty(self, doc):

        for i, p in enumerate(doc.paragraphs[:2]):
            if p.text and p.text in 'Sorry, your request could not be processed. The following error has occurred.\nDocument Number not specified for document retrieval':
                return True

        return False

    @classmethod
    def from_mapp_folder(cls, folder_path=None, get_parent=False, map={}):

        output = {'flags': []}
        if folder_path:
            files = os.listdir(folder_path)
        else:
            files = os.listdir()
        files.sort(reverse=True)

        all_files = []

        latest = {'Authorization': False, 'Amendment':False, 'Amendment_Authorization':False, 'Parallel':False}
        active_substances = {}

        for i, f in enumerate(files):
            if folder_path:
                file = cls(folder_path + f, convert=False if 'docx' in f else True)
            else:
                file = cls(f, convert=False if 'docx' in f else True)

            if f in map:
                file.document_url = map[f]

            try:
                doc = file.read_docx()
            except Exception as e:
                try:
                    doc = file.read_docx()
                except:
                    continue
            else:
                if file.is_empty(doc):
                    continue
            file.get_main_info(doc)

            if file.type == 'Authorization' and not latest['Authorization']:
                latest['Authorization'] = True

                if hasattr(file, 'formulation'):
                    formulation = file.formulation
                    output['formulation'] = formulation
                else:
                    output['flags'].append("No formulation available")

                protections = file.get_protections(doc)

                if hasattr(file, 'extent_of_authorization'):
                    output['extent_of_authorization'] = file.extent_of_authorization

                output['protections'] = {"protections":protections,
                                         "date": file.date_of_issue if hasattr(file, 'date_of_issue') else None,
                                         "file": file.path}

            elif file.type == 'Amendment':
                active_substances_list = file.get_active_substances(doc)
                for active_substance in active_substances_list:
                    if 'name' in active_substance:
                        if active_substance['name'] not in active_substances:
                            active_substances[active_substance['name']] = active_substance

                if not latest['Amendment']:
                    latest['Amendment'] = True

            elif file.type == 'Parallel' and get_parent and not latest['Parallel']:
                parent = file.get_parallel_parent(doc)
                output['parallel_parent'] = parent
                if parent['name'] != 'Not found':
                    latest['Parallel'] = True

            if (file.type == 'Amendment' or file.type == 'Authorization' or file.type == 'Parallel') and not latest['Amendment_Authorization']:

                crops, aquatic = file.get_crops(doc)
                if crops:
                    latest['Amendment_Authorization'] = True
                    output['crops'] = crops

                    if aquatic:
                        output['aquatic'] = aquatic

            output['flags'] += file.flags
            delattr(file, 'flags')

            if i == 0:
                output['expirations'] = file.expirations

            if hasattr(file, 'formulation'):
                delattr(file, 'formulation')
            if hasattr(file, 'extent_of_authorization'):
                delattr(file, 'extent_of_authorization')
            if hasattr(file, 'pesticides_exceptions'):
                delattr(file, 'pesticides_exceptions')
            if hasattr(file, 'extensions_exceptions'):
                delattr(file, 'extensions_exceptions')

            all_files.append(file.__dict__)

        if not output.get('crops'):
            output['flags'].append('No crops info available')
        #else:
            #output['crops'] = [dict(t) for t in {tuple(d.items()) for d in output['crops']}]

        if not output.get('protections') and output.get("Parallel Import") == "No":
            output['flags'].append('No protections info available (No Authorization file found)')
        elif output.get('protections'):
            if (output['protections']['protections']['Environmental protection:']['headers'][0]['value'] == "") and (
                    output['protections']['protections']['Other specific restrictions:']['headers'][0]['value'] == "") and (
                    output['protections']['protections']['Operator protection:']['headers'][0]['value'] == ""):
                output['flags'].append('No protections info available')

        if not output['flags']:
            del output['flags']

        output['files'] = all_files

        return output, active_substances

if __name__ == '__main__':
    docs, act_subs = DocumentFile.from_mapp_folder('Outputs/13-06-2020 12:08:48/pesticides/15923/', get_parent=True)
    docs['crops'].sort(key=lambda i: i['crop'])
    a = 2
    #file = DocumentFile('/home/xps/Documents/Dropbox/Freelancer/ABIERTO-MarkP/Daniel_Mark/pesticide-scraper-uk/Scripts/Outputs/04-02-2020 11:28:11/extensions/20082887.docx', True)
    #doc = file.read_docx()
    #protections = file.get_protections(doc)
    #crops, aquatic = file.get_crops(doc)
    '''
    import json
    pesticides_file = open('Outputs/26-02-2020 13:56:39/pesticide.json', 'r').read()
    pesticides = json.loads(pesticides_file)
    total = len(pesticides)
    for i, (mapp, p) in enumerate(pesticides.items()):
        docs, act_subs = DocumentFile.from_mapp_folder('Outputs/13-06-2020 12:08:48/pesticides/{}/'.format(mapp), get_parent=True)
        print(i + 1, '/', total)
    '''

